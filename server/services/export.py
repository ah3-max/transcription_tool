"""匯出渲染（S-09，FR-10/16/20/25）。

把產出內容（Markdown）渲染成 docx / md / txt 三格式（**不做 PDF**，NG-3／D-05）。
docx 套愛愛院字級／標題樣式（設計手冊 §7.2：H1 26 / H2 20 / 內文 11pt；H1 品牌主色＋粗體）。
內容約定為 Markdown：`# ` 文件標題、`## ` 章節、`- ` 項目、其餘為段落。
"""
import io
import re

MEDIA_TYPES = {
    "txt": "text/plain; charset=utf-8",
    "md": "text/markdown; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
ALLOWED_FMTS = set(MEDIA_TYPES)  # 僅 txt/md/docx；pdf 不在內（NG-3）

BRAND_PRIMARY = (0x00, 0xA9, 0x7A)  # #00A97A 品牌主色（§7.1）

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")


def _to_plain(content_md: str) -> str:
    """去除常見 Markdown 標記，回純文字（給 txt）。"""
    out = []
    for line in content_md.splitlines():
        m = _HEADING_RE.match(line)
        if m:
            out.append(m.group(2))
            continue
        b = _BULLET_RE.match(line)
        if b:
            out.append(f"・{b.group(1)}")
            continue
        out.append(line)
    return "\n".join(out)


def _to_docx(content_md: str) -> bytes:
    """以 python-docx 渲染；H1/H2 套樣式與品牌色（§7.2）。"""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    # 內文預設 11pt（§7.2）
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)

    for raw in content_md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        m = _HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            if level == 1:
                h = doc.add_heading(text, level=1)
                run = h.runs[0] if h.runs else h.add_run(text)
                run.bold = True
                run.font.size = Pt(26)
                run.font.color.rgb = RGBColor(*BRAND_PRIMARY)  # H1 品牌主色
            elif level == 2:
                h = doc.add_heading(text, level=2)
                run = h.runs[0] if h.runs else h.add_run(text)
                run.font.size = Pt(20)
            else:
                h = doc.add_heading(text, level=min(level, 9))
                if h.runs:
                    h.runs[0].font.size = Pt(16)
            continue
        b = _BULLET_RE.match(line)
        if b:
            doc.add_paragraph(b.group(1), style="List Bullet")
            continue
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render(content_md: str, fmt: str, *, title: str | None = None) -> tuple[bytes, str, str]:
    """渲染內容成指定格式。回 (bytes, media_type, ext)。

    fmt 須為 txt/md/docx；其他（含 pdf）→ ValueError（NG-3／D-05）。
    """
    fmt = (fmt or "").lower().lstrip(".")
    if fmt not in ALLOWED_FMTS:
        raise ValueError(f"不支援的格式：{fmt}（僅 docx/md/txt，不提供 PDF）")
    content_md = content_md or ""
    if title and not content_md.lstrip().startswith("#"):
        content_md = f"# {title}\n\n{content_md}"

    if fmt == "md":
        return content_md.encode("utf-8"), MEDIA_TYPES["md"], "md"
    if fmt == "txt":
        return _to_plain(content_md).encode("utf-8"), MEDIA_TYPES["txt"], "txt"
    return _to_docx(content_md), MEDIA_TYPES["docx"], "docx"
